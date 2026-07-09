#!/usr/bin/env python3
"""
Génère le dossier écrit du Bloc 3 CESIZen — « Déployer et sécuriser les applications
informatiques » (INFCDAAL3) au format Word éditable.

Sortie : ./Dossier_Bloc3_CESIZen.docx
Schémas : ./assets_bloc3/*.png (générés par generate_diagrams_bloc3.py)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Charte CESIZen ───────────────────────────────────────────────
BLUE    = RGBColor(0x00, 0x46, 0x82)
GREEN   = RGBColor(0x04, 0x96, 0x41)
GREEN_V = RGBColor(0x06, 0xC6, 0x56)
INK     = RGBColor(0x23, 0x23, 0x2D)
GREY    = RGBColor(0x6E, 0x6E, 0x6E)
RED     = RGBColor(0xC8, 0x10, 0x2E)
ORANGE  = RGBColor(0xE4, 0x80, 0x1C)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLUE_L  = "E6F0FA"
GREY_L  = "F2F2F2"

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets_bloc3")
OUT = os.path.join(HERE, "Dossier_Bloc3_CESIZen.docx")

doc = Document()

# ─── Styles de base ───────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color in [("Heading 1", 17, BLUE), ("Heading 2", 13.5, GREEN),
                          ("Heading 3", 11.5, BLUE)]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=INK, size=9.5, white=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = WHITE if white else color
    r.font.name = "Calibri"


def h1(text):
    doc.add_paragraph(text, style="Heading 1")


def h2(text):
    doc.add_paragraph(text, style="Heading 2")


def h3(text):
    doc.add_paragraph(text, style="Heading 3")


def para(text, size=10.5, italic=False, color=INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    r.font.color.rgb = color
    return p


def rich(segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, size=10.5):
    """segments = list of (text, bold)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    for text, bold in segments:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = INK
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    return p


def image(name, width=6.3, caption=None):
    path = os.path.join(ASSETS, name)
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Figure — " + caption)
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = GREY
        p.paragraph_format.space_after = Pt(8)


def callout(title, text, color=BLUE, fill=BLUE_L):
    """Encadré une-cellule."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    cell.width = Inches(6.4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def table(headers, rows, widths=None, header_fill="004682", zebra=True, size=9.3,
          fonts_bold_col0=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        _set_cell_text(hdr[i], htxt, bold=True, white=True, size=size,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            _set_cell_text(cells[ci], str(val), size=size,
                           bold=(fonts_bold_col0 and ci == 0))
            if zebra and ri % 2 == 1:
                _shade(cells[ci], GREY_L)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_break():
    doc.add_page_break()


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Faites un clic droit ici puis « Mettre à jour les champs » pour générer le sommaire."
    r.append(t)
    fldSimple.append(r)
    p._p.append(fldSimple)


# ══════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CESIZen"); r.font.size = Pt(46); r.font.bold = True; r.font.color.rgb = GREEN_V
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("L'application de votre santé mentale"); r.font.size = Pt(16); r.font.color.rgb = GREY
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Dossier — Activité 3"); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Déployer et sécuriser les applications informatiques")
r.font.size = Pt(15); r.italic = True; r.font.color.rgb = INK
for _ in range(6):
    doc.add_paragraph()
for line, sz, col, bold in [
    ("Plan de déploiement · Plan de maintenance · Plan de sécurisation", 12, GREEN, True),
    ("", 6, INK, False),
    ("Adam Marzuk — CESI École d'Ingénieurs", 12, INK, False),
    ("Titre Concepteur Développeur d'Applications (CDA)", 11, GREY, False),
    ("Bloc INFCDAAL3 — Commanditaire : Ministère des Solidarités et de la Santé", 11, GREY, False),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if line:
        r = p.add_run(line); r.font.size = Pt(sz); r.font.color.rgb = col; r.font.bold = bold
page_break()

# ══════════════════════════════════════════════════════════════════
# SOMMAIRE
# ══════════════════════════════════════════════════════════════════
h1("Sommaire")
add_toc()
page_break()

# ══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
h1("1. Introduction")
para("CESIZen est une plateforme grand public de santé mentale simulée pour le Ministère des "
     "Solidarités et de la Santé. Elle aide chacun à mieux comprendre sa santé mentale et à agir sur "
     "son stress à travers trois modules : la gestion des comptes utilisateurs, la diffusion "
     "d'informations et un tracker d'émotions. La solution repose sur une API REST Laravel, un site "
     "web Next.js et une application mobile React Native, adossés à une base PostgreSQL.")
para("Après la conception (Activité 1) et le développement/test du prototype (Activité 2), la présente "
     "Activité 3 traite du déploiement et de la sécurisation de l'application. Ce dossier présente les "
     "trois livrables attendus :")
bullet("l'architecture des environnements, le versioning et les automatisations de déploiement ;",
       bold_prefix="Un plan de déploiement — ")
bullet("un outil de ticketing configuré et la méthodologie de traitement des incidents et évolutions ;",
       bold_prefix="Un plan de maintenance — ")
bullet("l'analyse des vulnérabilités, les actions préventives/correctives, la protection des données "
       "personnelles et la gestion de crise.", bold_prefix="Un plan de sécurisation — ")
para("Les outils de versioning, d'intégration continue et de ticketing retenus (Gitea et Coolify) sont "
     "réellement configurés et font l'objet d'une démonstration lors de la soutenance.")

callout("Périmètre.", "Un environnement de déploiement (la production) est réellement mis en place et "
        "accessible en ligne sur cesizen.cleanows.fr. Les trois modules du prototype (Comptes, "
        "Informations, Tracker d'émotions) y sont déployés.")

# ══════════════════════════════════════════════════════════════════
# 2. PLAN DE DÉPLOIEMENT
# ══════════════════════════════════════════════════════════════════
h1("2. Plan de déploiement")

h2("2.1 Architecture technique de la solution")
para("L'application suit une architecture découplée : une API REST centrale expose la logique métier et "
     "les données, consommée par deux clients distincts (web et mobile). Ce découpage garantit la "
     "non-duplication de la logique et facilite l'ajout de nouveaux clients.")
table(
    ["Composant", "Technologie", "Rôle"],
    [
        ["API / Back-end", "Laravel 12 (PHP 8.4)", "API REST JSON, authentification JWT, logique métier"],
        ["Base de données", "PostgreSQL 16", "Persistance ACID des données"],
        ["Cache / file (dev)", "Redis 7", "Cache, sessions, files d'attente (environnement de dev)"],
        ["Front-end web", "Next.js 16 / React 19", "Site public et back-office, design system DSFR de l'État"],
        ["Application mobile", "React Native / Expo 54", "Application iOS & Android (Mobile First)"],
        ["Conteneurisation", "Docker / Docker Compose", "Isolation et reproductibilité des environnements"],
        ["Plateforme de déploiement", "Coolify (PaaS auto-hébergé)", "Build, orchestration, HTTPS, supervision"],
    ],
    widths=[1.7, 1.9, 2.8],
)
para("En production, le cache, les sessions et les files d'attente basculent sur la base de données "
     "(pilote « database »), ce qui évite d'exploiter Redis et simplifie l'infrastructure.", size=9.7,
     italic=True, color=GREY)

h2("2.2 Les trois environnements")
para("Le cycle de vie du logiciel s'appuie sur trois environnements distincts, du poste du développeur "
     "jusqu'à la production exposée aux utilisateurs.")
table(
    ["", "Développement", "Test / Intégration", "Production"],
    [
        ["Objectif", "Coder et tester en local", "Valider automatiquement chaque modification", "Servir les utilisateurs finaux"],
        ["Hébergement", "Poste du développeur", "Runner Gitea Actions (CI)", "Serveur Coolify (VPS)"],
        ["Orchestration", "docker-compose.yml", "Conteneur de CI éphémère", "docker-compose.prod.yaml"],
        ["Base de données", "PostgreSQL 16 + Redis", "SQLite en mémoire", "PostgreSQL 16 (réseau interne)"],
        ["Configuration", "APP_DEBUG=true, secrets en dur (non sensibles)", "APP_ENV=testing", "APP_DEBUG=false, secrets générés hors dépôt"],
        ["Accès", "localhost:3000 / :8000", "Journaux de la CI", "HTTPS : cesizen.cleanows.fr"],
    ],
    widths=[1.3, 1.85, 1.7, 1.9], size=8.8,
)

h2("2.3 Conteneurisation et images Docker")
para("Chaque service est livré sous forme d'image Docker, garantissant un fonctionnement identique "
     "d'un environnement à l'autre. Les images de développement privilégient le rechargement à chaud ; "
     "les images de production sont optimisées.")
bullet("l'image back-end (php:8.4-cli) installe les extensions PostgreSQL/bcmath et n'embarque que les "
       "dépendances de production (composer install --no-dev) ;", bold_prefix="Back-end — ")
bullet("l'image front-end utilise un build multi-étapes (builder → runner) et la sortie « standalone » "
       "de Next.js, produisant une image d'exécution minimale ;", bold_prefix="Front-end — ")
bullet("un script d'entrée (entrypoint.sh) automatise au démarrage la génération des clés (APP_KEY, "
       "JWT_SECRET), les migrations, le seed idempotent et la mise en cache des routes/vues.",
       bold_prefix="Démarrage — ")

h2("2.4 Automatisations et intégration continue")
para("Deux automatisations complémentaires structurent la chaîne de livraison :")
bullet("à chaque push ou pull request, le pipeline .gitea/workflows/ci.yml exécute les 29 tests "
       "automatisés (19 back-end PHPUnit, 6 web Vitest, 4 mobile Jest) et les contrôles de style "
       "(Laravel Pint, ESLint). Un changement dont les tests échouent ne peut pas être livré.",
       bold_prefix="Intégration continue (Gitea Actions) — ")
bullet("un merge sur la branche main déclenche, via un webhook, la reconstruction des images et le "
       "redéploiement par Coolify, avec migrations et seed joués automatiquement, health checks et "
       "redémarrage automatique des conteneurs.", bold_prefix="Déploiement continu (Coolify) — ")
image("archi_deploiement.png", width=6.4,
      caption="Chaîne de déploiement : trois environnements et flux d'intégration/livraison continues.")

h2("2.5 Étapes de déploiement et ressources")
para("Le déploiement en production se déroule selon les étapes suivantes, entièrement automatisées "
     "après le déclenchement :")
table(
    ["Étape", "Action", "Automatisation"],
    [
        ["1. Récupération", "Coolify clone la révision de main", "Webhook Gitea → Coolify"],
        ["2. Build", "Construction des images back & front", "Dockerfile.prod (multi-étapes)"],
        ["3. Secrets", "Génération APP_KEY / JWT_SECRET si absents", "entrypoint.sh"],
        ["4. Base de données", "Migrations + seed idempotent", "artisan migrate --force / db:seed"],
        ["5. Démarrage", "Lancement des conteneurs + health checks", "docker-compose.prod.yaml"],
        ["6. Exposition", "Routage HTTPS/TLS et certificat", "Reverse proxy Coolify"],
        ["7. Supervision", "Vérification /up et redémarrage auto", "healthcheck + restart: unless-stopped"],
    ],
    widths=[1.4, 2.6, 2.4], size=9,
)
para("Ressources et dimensionnement — l'application grand public est dimensionnée pour un VPS unique "
     "(2 vCPU, 4 Go de RAM, 40 Go SSD) hébergeant les trois conteneurs, suffisant pour la phase de "
     "test de 12 mois. La montée en charge se fait horizontalement (réplication des conteneurs back "
     "et front derrière le reverse proxy, base de données managée) sans refonte, l'API étant stateless.")
page_break()

# ══════════════════════════════════════════════════════════════════
# 3. VERSIONING
# ══════════════════════════════════════════════════════════════════
h1("3. Gestion des sources et versioning")
h2("3.1 Outil retenu : Gitea")
para("Le versioning s'appuie sur Git, hébergé sur une instance Gitea auto-hébergée. Gitea a été retenu "
     "car il réunit dans un seul outil, souverain et gratuit, le dépôt Git, le suivi des tickets, "
     "l'intégration continue (Gitea Actions) et la gestion des versions — cohérent avec un projet porté "
     "par un ministère (hébergement maîtrisé, pas de dépendance à un service tiers hors UE).")
table(
    ["Critère", "Gitea", "GitHub", "GitLab"],
    [
        ["Souveraineté / auto-hébergement", "Oui (léger)", "Non (SaaS US)", "Oui (lourd)"],
        ["Ticketing intégré", "Oui", "Oui", "Oui"],
        ["CI/CD intégrée", "Oui (Actions)", "Oui", "Oui"],
        ["Coût", "Gratuit", "Gratuit/payant", "Gratuit/payant"],
        ["Ressources serveur", "Faibles", "—", "Élevées"],
    ],
    widths=[2.4, 1.4, 1.4, 1.4], size=9,
)

h2("3.2 Stratégie de branches et conventions")
bullet("main : code en production, toujours déployable et protégé (merge via pull request + CI verte) ;")
bullet("develop : intégration des développements en cours ;")
bullet("feature/*, fix/*, hotfix/* : une branche par fonctionnalité ou correctif, fusionnée par PR.")
para("Les messages de commit suivent la convention Conventional Commits déjà appliquée sur le projet "
     "(feat, fix, docs, chore, refactor…) avec un scope (back, web, mobile, deploy), ce qui rend "
     "l'historique lisible et permet de générer automatiquement les notes de version. Les versions "
     "livrées sont marquées par des tags (v1.0, v1.1…) reliés aux jalons Gitea.")

# ══════════════════════════════════════════════════════════════════
# 4. PLAN DE MAINTENANCE
# ══════════════════════════════════════════════════════════════════
h1("4. Plan de maintenance")
h2("4.1 Outil de gestion des corrections et évolutions")
para("Le suivi des demandes s'appuie sur le module Issues de Gitea, complété par un tableau de bord "
     "Kanban (Projects) et des jalons (Milestones). Deux gabarits d'issues versionnés dans le dépôt "
     "structurent la saisie et posent automatiquement les bons libellés :")
bullet("gabarit incident (maintenance corrective) : sévérité, module, environnement, étapes de "
       "reproduction, résultat attendu/obtenu ;", bold_prefix="🐞 ")
bullet("gabarit demande d'évolution (maintenance évolutive) : besoin, solution proposée, priorité, "
       "critères d'acceptation.", bold_prefix="✨ ")
para("Les libellés couvrent le type, la gravité, la priorité, le statut et le module ; les jalons "
     "correspondent aux versions. Le tableau Kanban (À qualifier → À faire → En cours → En revue → "
     "Résolu) offre une vision temps réel de l'avancement, partagée entre le prestataire et le client.")

h2("4.2 Maintenance corrective : méthodologie et engagements de service")
para("La maintenance corrective démarre dès la prise en compte d'un ticket d'incident. La sévérité "
     "conditionne les délais contractuels de prise en compte et de correction, distinguant incident "
     "bloquant, majeur et mineur.")
table(
    ["Sévérité", "Définition", "Prise en compte", "Correction"],
    [
        ["Bloquant — critique", "Service inutilisable, aucun contournement", "1 h ouvrée", "3 h ouvrées"],
        ["Bloquant — fort", "Service inutilisable, contournement partiel", "2 h ouvrées", "6 h ouvrées"],
        ["Majeur", "Service dégradé / intermittent avec contournement", "7 h ouvrées", "16 h ouvrées"],
        ["Mineur (par lots)", "Altération mineure, service opérationnel", "1 j ouvré", "40 h ouvrées"],
    ],
    widths=[1.5, 2.5, 1.2, 1.2], size=9,
)
para("Le traitement suit un cycle standardisé et traçable, du signalement à la clôture validée par le "
     "client, en passant par la correction sur une branche dédiée et la revue par pull request "
     "(fusion conditionnée à une CI verte).")
image("cycle_ticket.png", width=6.4,
      caption="Cycle de vie d'un ticket sur Gitea et rappel des engagements de service (SLA).")

h2("4.3 Maintenance évolutive")
para("Chaque demande d'évolution fait l'objet d'un ticket puis d'une analyse par le prestataire "
     "proposant : la documentation fonctionnelle et technique, le délai de mise en œuvre et le coût "
     "associé. Après validation du client, l'évolution est planifiée dans un jalon de version. "
     "L'architecture découplée et le référentiel de données extensible permettent d'activer les "
     "modules écartés (Diagnostic de stress, Cohérence cardiaque, Activités de détente) sans réécrire "
     "le modèle existant.")

h2("4.4 Pilotage")
para("Le pilotage de la maintenance s'appuie sur les tableaux de bord natifs de Gitea : filtres par "
     "libellé/jalon/assigné, vue Kanban et suivi d'avancement par version. Ces indicateurs (nombre de "
     "tickets ouverts/fermés, respect des délais, répartition par module) permettent au Ministère de "
     "suivre la qualité de service.")

h2("4.5 Veille technologique et pérennité")
para("La pérennité de l'application dépend du suivi des évolutions de son écosystème. Une méthodologie "
     "de veille structurée est mise en place :")
table(
    ["Axe de veille", "Sources", "Rythme", "Action"],
    [
        ["Sécurité / vulnérabilités", "CVE, bulletins CERT-FR, GitHub Advisories, OWASP",
         "Continu (automatisé)", "Alertes de dépendances + correctifs prioritaires"],
        ["Frameworks (Laravel, Next.js, Expo)", "Blogs officiels, notes de version, roadmaps",
         "Mensuel", "Planifier les montées de version LTS"],
        ["Dépendances", "npm audit, composer audit, Renovate/Dependabot",
         "Hebdomadaire", "PR automatiques de mise à jour, testées par la CI"],
        ["Réglementation (RGPD, HDS)", "CNIL, ANSSI, Legifrance",
         "Trimestriel", "Adapter les mesures de conformité"],
    ],
    widths=[1.9, 2.3, 1.1, 1.9], size=8.7,
)
para("La veille sécurité est en partie automatisée : les outils d'analyse de dépendances (composer "
     "audit, npm audit) sont intégrables à la CI et signalent immédiatement toute dépendance "
     "vulnérable, tandis qu'un robot de mise à jour propose des pull requests validées par les tests.")
page_break()

# ══════════════════════════════════════════════════════════════════
# 5. PLAN DE SÉCURISATION
# ══════════════════════════════════════════════════════════════════
h1("5. Plan de sécurisation")
para("S'agissant d'une application portée par un ministère et manipulant des données de santé mentale "
     "(données sensibles au sens de l'article 9 du RGPD), la sécurité est traitée comme une exigence de "
     "premier ordre. Ce plan identifie les vulnérabilités et risques, en évalue la criticité, puis "
     "présente les mesures déjà en place et les actions correctives et préventives planifiées.")

h2("5.1 Analyse des vulnérabilités")
para("L'analyse s'appuie sur le référentiel OWASP Top 10. Le tableau ci-dessous confronte chaque risque "
     "aux mesures réellement implémentées dans le prototype et aux axes de renforcement.")
table(
    ["Vulnérabilité (OWASP)", "Mesure en place", "Renforcement prévu"],
    [
        ["A01 Contrôle d'accès défaillant", "RBAC par middleware (role:administrateur), garde-fous (un admin ne peut se désactiver/supprimer)", "Revue périodique des droits"],
        ["A02 Défaillances cryptographiques", "HTTPS/TLS, bcrypt (coût 12), mots de passe masqués et exclus des logs", "Chiffrement applicatif au repos des notes de santé (cast chiffré)"],
        ["A03 Injection", "ORM Eloquent (requêtes paramétrées), Form Requests de validation", "Purification/échappement HTML côté API"],
        ["A05 Mauvaise configuration", "APP_DEBUG=false en prod, secrets hors dépôt, CORS en liste blanche", "En-têtes de sécurité HSTS/CSP/X-Frame-Options"],
        ["A07 Authentification défaillante", "JWT signé (HS256), TTL 60 min, blacklist au logout, rate limiting", "Politique de mot de passe renforcée (complexité), vérification d'e-mail, MFA admin"],
        ["A09 Journalisation insuffisante", "Piste d'audit immuable (IP, ancienne/nouvelle valeur)", "Centralisation des logs + alerting (supervision)"],
    ],
    widths=[2.1, 2.6, 1.7], size=8.6,
)

h2("5.2 Matrice de criticité des risques")
para("Chaque risque est positionné selon sa gravité et sa probabilité ; la criticité est le produit des "
     "deux (échelle de 1 à 25). Cette hiérarchisation oriente la priorité des actions.")
image("matrice_risques.png", width=6.2,
      caption="Matrice de criticité des risques du projet CESIZen.")
table(
    ["Réf.", "Risque", "Gravité", "Prob.", "Criticité", "Priorité"],
    [
        ["R1", "Fuite de données de santé (art. 9)", "Critique (4)", "Faible (2)", "8", "Élevée"],
        ["R7", "Absence de chiffrement au repos", "Critique (4)", "Moyenne (3)", "12", "Élevée"],
        ["R2", "Injection SQL / XSS", "Grave (3)", "Moyenne (3)", "9", "Élevée"],
        ["R3", "Compromission d'un compte admin", "Grave (3)", "Faible (2)", "6", "Moyenne"],
        ["R6", "Dépendances vulnérables", "Modérée (2)", "Forte (4)", "8", "Moyenne"],
        ["R4", "DDoS / indisponibilité", "Modérée (2)", "Moyenne (3)", "6", "Moyenne"],
        ["R5", "Perte de données (sauvegarde)", "Modérée (2)", "Faible (2)", "4", "Faible"],
        ["R8", "Erreur humaine de déploiement", "Mineure (1)", "Faible (2)", "2", "Faible"],
    ],
    widths=[0.5, 2.5, 1.1, 1.0, 0.9, 0.9], size=8.7,
)

h2("5.3 Actions préventives et correctives")
para("Pour chaque risque, des actions préventives (réduire la probabilité) et correctives (limiter "
     "l'impact) sont définies, couvrant l'ensemble des risques critiques et majeurs identifiés.")
table(
    ["Réf.", "Actions préventives", "Actions correctives"],
    [
        ["R1", "Minimisation des données, chiffrement en transit, RBAC, audit des accès, hébergement UE/HDS",
         "Procédure de notification CNIL sous 72 h, isolation, révocation des accès, information des personnes"],
        ["R7", "Chiffrement applicatif des champs sensibles (note), chiffrement du volume PostgreSQL",
         "Rotation des clés, restauration depuis sauvegarde chiffrée"],
        ["R2", "Validation systématique (Form Requests), ORM paramétré, échappement des sorties",
         "Correctif prioritaire, revue de code, test de non-régression"],
        ["R3", "MFA administrateur, mots de passe robustes, rate limiting, journalisation des actions",
         "Blacklist du token, réinitialisation, analyse de la piste d'audit"],
        ["R6", "composer/npm audit en CI, robot de mise à jour, veille CVE",
         "Application du correctif, redéploiement automatisé"],
        ["R4", "Reverse proxy, rate limiting, health checks + redémarrage auto",
         "Mise à l'échelle, activation d'une protection anti-DDoS (CDN/WAF)"],
        ["R5", "Sauvegardes quotidiennes automatisées et testées, réplication du volume",
         "Restauration point-in-time, plan de reprise d'activité"],
    ],
    widths=[0.5, 3.0, 3.0], size=8.4,
)

h2("5.4 Chiffrement et protection des échanges")
bullet("l'ensemble des échanges entre les clients et l'API passe par HTTPS/TLS (terminaison au niveau "
       "du reverse proxy Coolify), interdisant l'interception en clair ;", bold_prefix="En transit — ")
bullet("les mots de passe sont hachés avec bcrypt (coût 12) et ne sont jamais renvoyés ni journalisés ; "
       "la clé applicative (APP_KEY, AES-256) protège cookies et données chiffrées. Le chiffrement "
       "applicatif des notes de santé (champ le plus sensible) est planifié via le cast chiffré de "
       "Laravel ;", bold_prefix="Au repos — ")
bullet("l'authentification repose sur des jetons JWT signés (HS256), d'une durée de vie de 60 minutes, "
       "invalidés à la déconnexion (liste noire).", bold_prefix="Jetons — ")

h2("5.5 Données personnelles et conformité RGPD")
para("Les données de santé mentale (émotions, intensité, notes) constituent des données sensibles "
     "(art. 9 RGPD). La conformité est prise en compte par conception :")
table(
    ["Exigence RGPD", "Mise en œuvre dans CESIZen"],
    [
        ["Base légale / consentement", "Consentement explicite obligatoire à l'inscription (non pré-coché), tracé en base"],
        ["Minimisation", "Seules les données nécessaires au service sont collectées"],
        ["Droit d'accès et portabilité (art. 15/20)", "Export JSON téléchargeable de toutes les données de l'utilisateur"],
        ["Droit à l'effacement (art. 17)", "Anonymisation (dépersonnalisation des notes de santé) et suppression logique (soft delete)"],
        ["Traçabilité / responsabilité", "Piste d'audit immuable (action, IP, anciennes/nouvelles valeurs, hors mot de passe)"],
        ["Localisation", "Hébergement dans l'Union européenne, aucun transfert hors UE ; cible HDS pour la prod"],
        ["Sécurité (art. 32)", "Chiffrement en transit, hachage, contrôle d'accès, journalisation"],
    ],
    widths=[2.3, 4.1], size=9,
)

h2("5.6 Gestion de crise et escalade d'information")
para("En cas d'incident de sécurité avéré (intrusion, fuite de données, indisponibilité majeure), une "
     "procédure de gestion de crise organise la réaction et la remontée d'information selon des niveaux "
     "d'escalade définis.")
table(
    ["Niveau", "Déclencheur", "Responsable", "Actions & communication"],
    [
        ["N1 — Détection", "Alerte supervision / signalement", "Astreinte technique", "Qualifier, ouvrir un ticket incident de sécurité, contenir"],
        ["N2 — Confirmation", "Incident de sécurité confirmé", "Responsable technique", "Isoler, préserver les preuves (logs), évaluer l'impact"],
        ["N3 — Crise", "Fuite de données personnelles avérée", "Chef de projet + DPO", "Notification CNIL sous 72 h, information des personnes concernées"],
        ["N4 — Direction", "Impact majeur / médiatique", "Direction + Ministère", "Communication officielle, coordination ANSSI si nécessaire"],
    ],
    widths=[1.3, 1.9, 1.4, 1.8], size=8.6,
)
para("La communication distingue la remontée interne (technique → chef de projet → direction) et la "
     "communication externe (personnes concernées, CNIL sous 72 h conformément à l'art. 33 RGPD, "
     "éventuellement ANSSI). Chaque crise donne lieu à un retour d'expérience et à des actions "
     "préventives pour éviter la récurrence.")
page_break()

# ══════════════════════════════════════════════════════════════════
# 6. BONNES PRATIQUES
# ══════════════════════════════════════════════════════════════════
h1("6. Bonnes pratiques de développement")
para("La maintenabilité et la sécurité de l'application reposent sur des pratiques de développement "
     "structurées, appliquées tout au long du projet :")
bullet("architecture MVC côté back-end (Modèles, Contrôleurs fins, Form Requests de validation, "
       "couche Service pour les rapports) ; App Router, hooks personnalisés et store côté front.",
       bold_prefix="Structure du code — ")
bullet("respect du PSR-12 vérifié par Laravel Pint ; TypeScript strict et ESLint côté web et mobile ; "
       "contrôles intégrés à la CI.", bold_prefix="Normes et style — ")
bullet("29 tests automatisés sur les trois couches (PHPUnit, Vitest, Jest), exécutés à chaque push, "
       "empêchant toute régression de partir en production.", bold_prefix="Tests — ")
bullet("Conventional Commits, revue systématique par pull request, historique lisible.",
       bold_prefix="Gestion de versions — ")
bullet("README, guide de lancement, documentation d'API commentée et guide de migration/outillage "
       "(MIGRATION_GITEA.md).", bold_prefix="Documentation — ")

h1("7. Conclusion")
para("Le projet CESIZen dispose d'une chaîne de déploiement complète et outillée : trois environnements "
     "clairement séparés, un versioning souverain sous Gitea, une intégration continue exécutant "
     "29 tests à chaque modification et un déploiement continu automatisé vers une production réellement "
     "en ligne et sécurisée par HTTPS. La maintenance est cadrée par un outil de ticketing configuré, "
     "des engagements de service précis et une veille technologique méthodique. La sécurisation, enfin, "
     "repose sur une analyse de risques hiérarchisée, des mesures déjà implémentées (authentification "
     "JWT, RBAC, audit, RGPD) et un plan d'actions correctives et préventives, complété par une "
     "procédure de gestion de crise. L'application est ainsi prête à être exploitée et à évoluer dans "
     "la durée au service de la santé mentale du grand public.")

# ─── Pied de page ─────────────────────────────────────────────────
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("CESIZen · Bloc INFCDAAL3 — Déployer et sécuriser les applications informatiques · Adam Marzuk")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY

doc.save(OUT)
print("Dossier généré :", OUT)
print("Paragraphes :", len(doc.paragraphs), "· Tableaux :", len(doc.tables))
